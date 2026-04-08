# agents/e_output.py
import json
import os
from datetime import datetime
from docx import Document
from state import DirectorState

OUTPUT_BASE = os.path.join(os.path.dirname(__file__), "..", "output")

def _build_json(state: DirectorState) -> dict:
    script = state["script"]
    scenes_data = state["scenes"].get("scenes", [])
    style_scenes = {s["scene_id"]: s for s in state["style"].get("scenes", [])}

    characters = [
        {
            "name": c["name"],
            "prompt": (
                f"{c.get('appearance', '')}, {c.get('costume', '')}, "
                f"状态：{c.get('state', '')}"
            )
        }
        for c in script.get("characters", [])
    ]

    scenes = []
    for scene in scenes_data:
        sid = scene["scene_id"]
        style = style_scenes.get(sid, {})
        scenes.append({
            "scene_id": sid,
            "description": scene.get("location", ""),
            "prompt": (
                f"{scene.get('location', '')}, {scene.get('background', '')}, "
                f"{scene.get('time', '')}, {scene.get('weather', '')}, "
                f"{scene.get('lighting', '')}, 镜头：{style.get('camera', '')}"
            ),
            "editing_technique": style.get("editing_technique", ""),
            "filter_style": style.get("filter_style", ""),
        })

    return {
        "project": state["user_input"][:20],
        "global_style_prompt": state["style"].get("global_style", ""),
        "characters": characters,
        "scenes": scenes,
    }

def _build_md(state: DirectorState, data: dict) -> str:
    script = state["script"]
    script_scenes = {s["scene_id"]: s for s in script.get("scenes", [])}
    lines = [
        f"# 分镜文档\n",
        f"**故事描述**: {state['user_input']}\n",
        f"**故事类型**: {state['story'].get('story_type')}  ",
        f"**情感基调**: {state['story'].get('tone')}  ",
        f"**核心冲突**: {state['story'].get('conflict')}\n",
        f"**整体风格**: {data['global_style_prompt']}\n",
        "---\n",
        "## 角色一览\n",
    ]
    for c in script.get("characters", []):
        lines.append(f"### {c['name']}")
        lines.append(f"- **背景**: {c.get('background', '')}")
        lines.append(f"- **外貌**: {c.get('appearance', '')}")
        lines.append(f"- **服装**: {c.get('costume', '')}")
        lines.append(f"- **状态**: {c.get('state', '')}\n")

    lines.append("---\n")
    lines.append("## 场景分镜\n")
    for scene_data in data["scenes"]:
        sid = scene_data["scene_id"]
        script_scene = script_scenes.get(sid, {})
        lines.append(f"### 场景 {sid}：{script_scene.get('title', '')}")
        lines.append(f"**地点**: {scene_data['description']}  ")
        lines.append(f"**AI提示词**: `{scene_data['prompt']}`\n")
        lines.append("**对白**:")
        for d in script_scene.get("dialogue", []):
            lines.append(f"> {d['speaker']}：{d['line']}")
        lines.append(f"\n**拍摄建议**: {style_hint(scene_data['prompt'])}")
        lines.append(f"**剪辑手法**: {scene_data['editing_technique']}")
        lines.append(f"**推荐滤镜**: {scene_data['filter_style']}\n")
        lines.append("---\n")
    return "\n".join(lines)

def style_hint(prompt: str) -> str:
    return prompt.split("镜头：")[-1] if "镜头：" in prompt else ""

def _build_docx(md_text: str, docx_path: str):
    doc = Document()
    for line in md_text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(docx_path)

def output_node(state: DirectorState) -> dict:
    project_name = datetime.now().strftime("%Y%m%d_%H%M%S")
    project_dir = os.path.join(OUTPUT_BASE, project_name)
    os.makedirs(project_dir, exist_ok=True)

    data = _build_json(state)
    json_path = os.path.join(project_dir, "prompts.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    md_text = _build_md(state, data)
    md_path = os.path.join(project_dir, "storyboard.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_text)

    docx_path = os.path.join(project_dir, "storyboard.docx")
    _build_docx(md_text, docx_path)

    return {"output": {"json_path": json_path, "md_path": md_path, "docx_path": docx_path}}
